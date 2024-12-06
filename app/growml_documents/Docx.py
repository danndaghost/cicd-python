import os
import shutil
import subprocess
import tempfile
from io import BytesIO
from pathlib import Path
from zipfile import BadZipFile

import docx
import numpy as np
import pandas as pd
import pymupdf
from docx.table import Table
from docx.text.paragraph import Paragraph

from growml_documents.Document import Document


class Docx(Document):
    def __init__(self, filepath_or_buffer: str | Path | BytesIO):
        super().__init__(filepath_or_buffer)
        try:
            self.docx = docx.Document(self.file_path)
        except BadZipFile:
            self.docx = docx.Document()
        self.temp_dir = tempfile.mkdtemp()

    def read_text(self):
        text = []
        for e in self.docx.iter_inner_content():
            if isinstance(e, Paragraph):
                text.append(f"{e.text}\n")
            elif isinstance(e, Table):
                for row in e.table.rows:
                    for cell in row.cells:
                        text.append(f"{cell.text}\t")
                    text.append("\n")
        return "".join(text)

    def _convert_to_pdf(self):
        # Prepare the command to convert DOCX to PDF
        command = [
            "libreoffice",
            "--headless",
            "--convert-to",
            "pdf",
            str(self.file_path),
            "--outdir",
            self.temp_dir,
        ]

        try:
            # Run the command
            subprocess.run(command, check=True)
            pdf_filename = (
                os.path.splitext(os.path.basename(self.file_path))[0] + ".pdf"
            )
            pdf_path = os.path.join(self.temp_dir, pdf_filename)
            print(f"Conversion successful! PDF saved to {pdf_path}")
            return pdf_path
        except subprocess.CalledProcessError as e:
            print(f"Error during conversion: {e}")
            return None

    def count_pages(self):
        try:
            pdf_path = self._convert_to_pdf()
            if pdf_path:
                pdf = pymupdf.open(pdf_path)
                num_pages = pdf.page_count
                pdf.close()
                return num_pages
            return 0
        finally:
            shutil.rmtree(self.temp_dir, ignore_errors=True)

    def get_paragraphs(self):
        return [p.text for p in self.docx.paragraphs]

    def get_tables(self):
        return [self._table_to_pandas(table) for table in self.docx.tables]

    def _table_to_pandas(self, table):
        n_rows = len(table.rows)
        n_cols = len(table.columns)
        content = np.empty((n_rows, n_cols), dtype=object)
        for i in range(n_rows):
            for j in range(n_cols):
                content[i][j] = table.cell(i, j).text
        df = pd.DataFrame(content)
        return df
